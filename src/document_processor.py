from typing import List, Dict, Optional
import re
import sys
from langchain_text_splitters import RecursiveCharacterTextSplitter
from src.logger import logging
from src.exception import CustomException

import PyPDF2
import pdfplumber
import fitz
from docx import Document
import magic
from pathlib import Path

class DocumentProcessor:
    """
        Enhanced document processor that supports docx,pdf,txt 
        Files with intelligent format detection and parsing 
    """
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]  # Added ". " for better sentence splitting
        )
        self.supported_formats={
            ".txt":self._parse_txt,
            ".pdf":self._parse_pdf,
            ".docx":self._parse_docx,
            ".doc":self._parse_docx
        }
        
    def detect_file_type(self,file_path:str):
        """
        Detect the file type of a given file using magic numbers (MIME type detection).
        This method attempts to identify the file type by reading the file's magic bytes
        and mapping the detected MIME type to a standard file extension. If the MIME type
        is not recognized in the mapping, or if magic detection fails, it falls back to
        using the file's extension.
        Args:
            file_path (str): The absolute or relative path to the file to be analyzed.
        Returns:
            str: The detected file extension (including the dot), such as ".txt", ".pdf", 
                 ".docx", or ".doc". If the MIME type cannot be mapped, returns the file's 
                 original extension in lowercase.
        Raises:
            No exceptions are raised; errors are caught and logged as warnings, with 
            fallback behavior to the file extension.
        Note:
            - Requires the 'python-magic' library to be installed.
            - Supported MIME types: text/plain, application/pdf, 
              application/vnd.openxmlformats-officedocument.wordprocessingml.document,
              application/word
            - For unmapped MIME types or when magic detection fails, the method returns 
              the file's original extension in lowercase.
        
        """    
        try:
            mime=magic.Magic(mime=True)
            file_type=mime.from_file(file_path)
            
            # Map MIME types to extensions
            mime_to_extensions={
                'text/plain':".txt",
                'application/pdf':".pdf",
                'application/vnd.openxlmformats-officedocument.wordprocessingml.document':".docx",
                'application/word':".doc"
            }
            
            detected_ext=mime_to_extensions.get(file_type)
            if detected_ext:
                logging.info(f"Detected file type:{file_type} -> {detected_ext}")
                return detected_ext
            
            # Fallback to extension
            return Path(file_path).suffix.lower() 
        except Exception as e:
            logging.warning(f"Magic detection failed, using extension: {e}")
            return Path(file_path).suffix.lower()
        
    def _parse_txt(self, file_path: str) -> str:
        """Parse plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            logging.info(f"✅ Parsed TXT: {len(text)} characters")
            return text
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    logging.info(f"✅ Parsed TXT with {encoding}: {len(text)} characters")
                    return text
                except:
                    continue
            
            raise CustomException(sys, "Could not decode text file with any encoding")    
        
    def _parse_pdf(self, file_path: str) -> str:
        """
        Parse PDF with multiple strategies:
        1. pdfplumber (best for tables and layout)
        2. PyMuPDF (fast, good for scanned docs)
        3. PyPDF2 (fallback)
        """
        try:
            # Strategy 1: pdfplumber (best quality)
            text = self._parse_pdf_pdfplumber(file_path)
            if text and len(text) > 100:
                logging.info(f"✅ Parsed PDF with pdfplumber: {len(text)} characters")
                return text
            
            # Strategy 2: PyMuPDF (faster)
            text = self._parse_pdf_pymupdf(file_path)
            if text and len(text) > 100:
                logging.info(f"✅ Parsed PDF with PyMuPDF: {len(text)} characters")
                return text
            
            # Strategy 3: PyPDF2 (fallback)
            text = self._parse_pdf_pypdf2(file_path)
            logging.info(f"✅ Parsed PDF with PyPDF2: {len(text)} characters")
            return text
            
        except Exception as e:
            logging.error(f"Error parsing PDF: {str(e)}")
            raise CustomException(sys, e)
    
    def _parse_pdf_pdfplumber(self, file_path: str) -> str:
        """Parse PDF using pdfplumber (best for tables)"""
        try:
            text_parts = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    
                    if page_text:
                        text_parts.append(f"\n--- Page {page_num} ---\n")
                        text_parts.append(page_text)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        text_parts.append(f"\n[Table {table_num}]\n")
                        for row in table:
                            text_parts.append(" | ".join([cell or "" for cell in row]))
                            text_parts.append("\n")
            
            return "".join(text_parts)
            
        except Exception as e:
            logging.warning(f"pdfplumber parsing failed: {e}")
            return ""
    
    def _parse_pdf_pymupdf(self, file_path: str) -> str:
        """Parse PDF using PyMuPDF (fast)"""
        try:
            text_parts = []
            
            with fitz.open(file_path) as doc:
                for page_num, page in enumerate(doc, 1):
                    text_parts.append(f"\n--- Page {page_num} ---\n")
                    text_parts.append(page.get_text())
            
            return "".join(text_parts)
            
        except Exception as e:
            logging.warning(f"PyMuPDF parsing failed: {e}")
            return ""
    
    def _parse_pdf_pypdf2(self, file_path: str) -> str:
        """Parse PDF using PyPDF2 (fallback)"""
        try:
            text_parts = []
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(reader.pages, 1):
                    text_parts.append(f"\n--- Page {page_num} ---\n")
                    text_parts.append(page.extract_text())
            
            return "".join(text_parts)
            
        except Exception as e:
            logging.warning(f"PyPDF2 parsing failed: {e}")
            return ""
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    text_parts.append("\n")
            
            # Extract tables
            for table_num, table in enumerate(doc.tables, 1):
                text_parts.append(f"\n[Table {table_num}]\n")
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_parts.append(row_text)
                    text_parts.append("\n")
            
            text = "".join(text_parts)
            logging.info(f"✅ Parsed DOCX: {len(text)} characters")
            return text
            
        except Exception as e:
            logging.error(f"Error parsing DOCX: {str(e)}")
            raise CustomException(sys, e)    
            
        
    def load_documents(self, file_path: str) -> str:
        """
            Load document with automatic format detection
            
            Args:
                file_path: Path to document file
                
            Returns:
                str: Extracted text content
        """
        try:
            #validate if file exists
            if not Path(file_path).exists():
                raise FileNotFoundError(f"File not found at: {file_path}")
            
            # Detect file type
            file_ext=self.detect_file_type(file_path)
            
            # Get appropriate parser
            parser=self.supported_formats.get(file_ext)
            if not parser:
                raise ValueError(f"Unsupported format: {file_ext}")
            
            # Parse document
            logging.info(f"Parsing the document:{file_path} (type:{file_ext})")
            text=parser(file_path)
            
            if not text or len(text) < 50:
                raise ValueError(f"Extracted text too short or empty: {len(text)}")
            
            text=self._clean_text(text)
            
            logging.info(f"SUccessfully loaded document:{len(text)} characters")
            
            return text
            
            # with open(file=file_path, mode='r', encoding='utf-8') as file:
            #     text = file.read()
            # logging.info(f"Successfully loaded document from {file_path}")
            # return text    
        except FileNotFoundError:
            logging.error(f"File not found: {file_path}")
            raise CustomException(sys, f"File not found: {file_path}")
        except Exception as e:
            logging.error(f"Error loading document: {str(e)}")
            raise CustomException(sys, e) 
        
    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove page markers if too many
        text = re.sub(r'\n--- Page \d+ ---\n', '\n', text)
        
        # Remove null characters
        text = text.replace('\x00', '')
        
        return text.strip()    
            
    def extract_sections(self, text: str) -> List[Dict]:
        """Extract sections from policy document"""
        try:
            # Fixed pattern - removed invalid regex syntax
            # This pattern captures from "SECTION X:" until next section or end
            pattern = r'SECTION\s+(\d+):\s+([A-Z\s]+)\s*\n(.*?)(?=SECTION\s+\d+:|$)'
            matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
            
            sections = []
            for i, match in enumerate(matches):
                section_number = match[0]  # First capture group
                title = match[1].strip()    # Second capture group
                content = match[2].strip()  # Third capture group
                
                sections.append({
                    "section_number": section_number,
                    "title": title,
                    "content": content
                })
                logging.info(f"Extracted Section {section_number}: {title}")
                logging.info(f"Content preview: {content[:100]}...")
            
            return sections
        except Exception as e:
            logging.error(f"Error extracting sections: {str(e)}")
            raise CustomException(sys, e)       
    
    def chunk_document(self, text: str) -> List[Dict]:
        """
        Split document into chunks with metadata.
        
        Returns:
            List[Dict]: Each dict contains:
                - id: unique identifier (chunk_0, chunk_1, etc.)
                - text: chunk content
                - section: section title/name
                - chunk_index: position in document
                - char_start: starting character position
                - char_end: ending character position
        """
        try:
            # Step 1: Split text into chunks
            split_texts = self.text_splitter.split_text(text)
            logging.info(f"Document split into {len(split_texts)} chunks")
            
            chunks = []
            current_position = 0
            
            # Step 2: Add metadata to each chunk
            for idx, chunk_text in enumerate(split_texts):
                # Find the position of this chunk in the original text
                chunk_start = text.find(chunk_text, current_position)
                
                # Handle case where exact match not found (due to processing)
                if chunk_start == -1:
                    chunk_start = current_position
                
                chunk_end = chunk_start + len(chunk_text)
                
                # Step 3: Identify which section this chunk belongs to
                section_name = self._identify_section(chunk_text, text, chunk_start)
                
                # Create chunk dictionary with metadata
                chunk_dict = {
                    "id": f"chunk_{idx}",
                    "text": chunk_text,
                    "section": section_name,
                    "chunk_index": idx,
                    "char_start": chunk_start,
                    "char_end": chunk_end
                }
                
                chunks.append(chunk_dict)
                current_position = chunk_end
                
                logging.info(f"Created chunk {idx}: section='{section_name}', length={len(chunk_text)}")
            
            return chunks
            
        except Exception as e:
            logging.error(f"Error chunking document: {str(e)}")
            raise CustomException(sys, e)
        
    def _identify_section(self, chunk: str, full_text: str, chunk_start: int) -> str:
        """
        Identify which section a chunk belongs to based on its position in the full text.
        
        Logic:
        1. Find all section headers and their positions
        2. Determine which section the chunk falls under based on character position
        3. Return the section title
        
        Args:
            chunk: The text chunk to identify
            full_text: The complete document text
            chunk_start: Starting position of chunk in full_text
            
        Returns:
            str: Section title or "General" if no section found
        """
        try:
            # Find all section headers with their positions
            pattern = r'SECTION\s+(\d+):\s+([A-Z\s]+)'
            matches = re.finditer(pattern, full_text, re.IGNORECASE)
            
            # Store sections with their start positions
            sections_positions = []
            for match in matches:
                section_number = match.group(1)
                section_title = match.group(2).strip()
                section_start = match.start()
                
                sections_positions.append({
                    'number': section_number,
                    'title': section_title,
                    'start': section_start
                })
            
            # Sort by position (should already be sorted, but just in case)
            sections_positions.sort(key=lambda x: x['start'])
            
            # Find which section this chunk belongs to
            # The chunk belongs to the last section that starts before the chunk
            current_section = "General"  # Default if no section found
            
            for section in sections_positions:
                if section['start'] <= chunk_start:
                    current_section = section['title']
                else:
                    # We've passed the chunk's position
                    break
            
            return current_section
            
        except Exception as e:
            logging.error(f"Error identifying section: {str(e)}")
            # Return default section name instead of failing
            return "General"
            
    def validate_chunks(self, chunks: List[Dict]) -> bool:
        """
        Validate that chunks have required metadata and reasonable content.
        
        Checks:
        - All required fields present
        - No empty chunks
        - Reasonable length distribution
        """
        try:
            if not chunks:
                logging.error("No chunks to validate")
                return False
            
            required_keys = ["id", "text", "section", "chunk_index", "char_start", "char_end"]
            
            for i, chunk in enumerate(chunks):
                # Check all required keys present
                if not all(key in chunk for key in required_keys):
                    missing_keys = [key for key in required_keys if key not in chunk]
                    logging.error(f"Chunk {i} missing metadata: {missing_keys}")
                    return False
                
                # Check chunk is not empty
                if not chunk['text'] or len(chunk['text'].strip()) == 0:
                    logging.error(f"Chunk {i} has empty text")
                    return False
                
                # Check chunk has reasonable length (not too short)
                if len(chunk['text']) < 50:
                    logging.warning(f"Chunk {i} is very short: {len(chunk['text'])} characters")
                
                # Check char_end > char_start
                if chunk['char_end'] <= chunk['char_start']:
                    logging.error(f"Chunk {i} has invalid character positions")
                    return False
            
            # Check chunk indices are sequential
            for i, chunk in enumerate(chunks):
                if chunk['chunk_index'] != i:
                    logging.error(f"Chunk indices not sequential at position {i}")
                    return False
            
            logging.info(f"All {len(chunks)} chunks validated successfully.")
            logging.info(f"Average chunk length: {sum(len(c['text']) for c in chunks) / len(chunks):.0f} characters")
            
            return True
            
        except Exception as e:
            logging.error(f"Error validating chunks: {str(e)}")
            raise CustomException(sys, e)
    
    def get_chunk_statistics(self, chunks: List[Dict]) -> Dict:
        """
        Get statistics about the chunks for quality assessment.
        
        Returns:
            Dict with statistics: count, avg_length, sections_covered, etc.
        """
        try:
            if not chunks:
                return {}
            
            chunk_lengths = [len(chunk['text']) for chunk in chunks]
            sections = set(chunk['section'] for chunk in chunks)
            
            stats = {
                'total_chunks': len(chunks),
                'avg_chunk_length': sum(chunk_lengths) / len(chunks),
                'min_chunk_length': min(chunk_lengths),
                'max_chunk_length': max(chunk_lengths),
                'unique_sections': len(sections),
                'sections_covered': list(sections)
            }
            
            logging.info(f"Chunk statistics: {stats}")
            return stats
            
        except Exception as e:
            logging.error(f"Error calculating statistics: {str(e)}")
            return {}